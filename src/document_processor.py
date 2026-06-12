from pathlib import Path
from typing import List
import logging

from langchain_community.document_loaders import PyPDFLoader, WebBaseLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

from src.settings import CHUNK_SIZE, CHUNK_OVERLAP, MAX_FILE_SIZE

logger = logging.getLogger(__name__)


def load_pdf(file_path: str) -> List[Document]:
    """
    Load a PDF from disk.
    Return a list of Document objects — one per page.
    Each Document has two things:
      - page_content: the raw text of that page
      - metadata: dict with info like {"page": 3, "source": "paper.pdf"}
    """
    try:
        loader = PyPDFLoader(file_path)
        pages = loader.load()
        return pages
    except Exception as e:
        msg = str(e).lower()
        if "encrypted" in msg:
            logger.warning("Encrypted PDF attempted: %s", file_path)
            raise ValueError("The PDF is encrypted and cannot be read.")
        if "cannot read" in msg or "no such file" in msg:
            logger.warning("Unreadable file: %s", file_path)
            raise ValueError(f"Cannot read the file: {file_path}")
        logger.exception("Failed to load PDF: %s", file_path)
        raise ValueError(f"Failed to load PDF: {e}")


def load_text(file_path: str) -> List[Document]:
    """
    Load a plain-text file (.txt) from disk.
    Return a single Document with the full file content.
    """
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        filename = Path(file_path).name
        return [Document(page_content=text, metadata={"source": filename, "page": 1})]
    except UnicodeDecodeError:
        raise ValueError("The text file is not valid UTF-8.")


def load_docx(file_path: str) -> List[Document]:
    """
    Load a Word document (.docx) from disk.
    Return a single Document with all paragraph text joined.
    """
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)
        filename = Path(file_path).name
        if not text.strip():
            raise ValueError("The Word document appears to be empty.")
        return [Document(page_content=text, metadata={"source": filename, "page": 1})]
    except ImportError:
        raise ValueError(
            "python-docx is required to process .docx files "
            "(pip install python-docx)."
        )


def load_url(url: str) -> List[Document]:
    """
    Load text content from a URL.
    Return a Document with the page content and URL as source.
    """
    try:
        loader = WebBaseLoader(url)
        docs = loader.load()
        for doc in docs:
            doc.metadata["source"] = url
            doc.metadata["page"] = 1
        return docs
    except ImportError:
        raise ValueError(
            "langchain_community is required to load URLs "
            "(pip install langchain-community)."
        )
    except Exception as e:
        logger.exception("Failed to load URL: %s", url)
        raise ValueError(f"Failed to load URL: {e}")


def chunk_documents(
    documents: List[Document],
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
) -> List[Document]:
    """
    Split documents into smaller overlapping chunks.

    RecursiveCharacterTextSplitter tries to split on:
      1. Paragraphs (\n\n)  ← preferred, keeps ideas together
      2. Lines (\n)
      3. Sentences (". ")
      4. Words (" ")
      5. Characters ("")    ← last resort
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_documents(documents)
    return chunks


def process_uploaded_file(uploaded_file, save_dir: str = "data", session_id: str = "") -> List[Document]:
    """
    Accept a Streamlit UploadedFile, save it to disk, then load and chunk it.
    If session_id is provided, files are stored in a session-scoped subdirectory.
    Supports: .pdf, .txt, .docx
    """
    if uploaded_file.size > MAX_FILE_SIZE:
        raise ValueError(
            f"File size ({uploaded_file.size / 1024 / 1024:.1f} MB) exceeds "
            f"the maximum allowed size of {MAX_FILE_SIZE / 1024 / 1024:.0f} MB."
        )

    save_path = Path(save_dir)
    if session_id:
        save_path = save_path / session_id
    save_path.mkdir(parents=True, exist_ok=True)

    file_path = save_path / Path(uploaded_file.name).name
    raw_bytes = uploaded_file.read()

    ext = Path(uploaded_file.name).suffix.lower()
    if ext == ".pdf":
        if raw_bytes[:4] != b"%PDF":
            raise ValueError("The file is not a valid PDF (missing %PDF header).")
    elif ext == ".txt":
        pass  # no magic-byte check needed
    elif ext == ".docx":
        if raw_bytes[:2] != b"PK":
            raise ValueError("The file is not a valid Word document (missing ZIP header).")
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    file_path.write_bytes(raw_bytes)

    try:
        if ext == ".pdf":
            pages = load_pdf(str(file_path))
        elif ext == ".txt":
            pages = load_text(str(file_path))
        elif ext == ".docx":
            pages = load_docx(str(file_path))
        else:
            raise ValueError(f"Unsupported file type: {ext}")
        chunks = chunk_documents(pages)
    except Exception:
        logger.exception("Failed to process uploaded file, cleaning up")
        file_path.unlink()
        raise

    return chunks
